import json
import logging
import itertools
from unicodedata import normalize

from django.apps import apps
from django.http import HttpResponse, JsonResponse,Http404
from django.conf import settings
from django.db.models import Prefetch
from django.core.mail import send_mail
from django.contrib.auth import get_user_model
from django.template.loader import render_to_string

from rest_framework import viewsets, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.exceptions import NotFound
from rest_framework.permissions import IsAdminUser

from fpdf import FPDF, HTMLMixin
from docx import Document
from docx.shared import Pt

from common.views import (
    CsrfExemptSessionAuthentication,
    StandardResultsSetPagination,
)
from .models import Request, FileRequest
from .serializers import RequestSerializer, RequestFileSerializer
import os
User = get_user_model()
Library = apps.get_model('library', 'Library')
Sample = apps.get_model('sample', 'Sample')
LibraryPreparation = apps.get_model(
    'library_preparation', 'LibraryPreparation')

logger = logging.getLogger('db')


class PDF(FPDF):  # pragma: no cover
    def __init__(self, title='Title', font='Arial'):
        self.title = title
        self.font = font
        super().__init__()

    def header(self):
        self.set_font(self.font, style='B', size=14)  # Arial bold 15
        self.cell(0, 10, self.title, align='C')       # Title
        self.ln(10)                                   # Line break

    def footer(self):
        self.set_y(-15)  # Position at 1.5 cm from bottom
        self.set_font(self.font, size=8)  # Arial 8
        # Page number
        self.cell(0, 10, 'Page ' + str(self.page_no()) + ' of {nb}', 0, 0, 'C')

    def info_row(self, title, value):
        self.set_font(self.font, style='B', size=11)
        self.cell(35, 10, title + ':')
        self.set_font(self.font, size=11)
        self.cell(0, 10, value)
        self.ln(6)

    def multi_info_row(self, title, value):
        self.set_font(self.font, style='B', size=11)
        self.ln(3)
        self.cell(35, 4, title + ':')
        self.set_font(self.font, size=11)
        self.multi_cell(0, 5, value)
        self.ln(6)

    def table_row(self, index, name, barcode, type, depth, bold=False):
        if bold:
            self.set_font(self.font, style='B', size=11)
        else:
            self.set_font(self.font, size=11)
        self.cell(10, 10, str(index))
        self.cell(60, 10, name)
        self.cell(40, 10, barcode)
        self.cell(35, 10, type)
        self.cell(0, 10, str(depth))
        self.ln(6)


class Report(FPDF, HTMLMixin):
    def __init__(self, title='Report', font='Arial'):
        self.title = title
        self.font = font
        super().__init__()

    def header(self):
        self.set_font(family=self.font, size=8)
        self.set_text_color(r=189, g=189, b=189)
        self.cell(0, 10, 'COMPLETE REPORT', align='L')
        self.cell(
            0, 10, 'Deep Sequencing Facility @ MPI-IE, Freiburg', align='R')
        self.ln(10)

    def footer(self):
        self.set_y(-15)  # Position at 1.5 cm from bottom
        self.set_font(self.font, size=8)  # Arial 8
        # Page number
        self.cell(0, 10, 'Page ' + str(self.page_no()) + ' of {nb}', 0, 0, 'C')

    def page_header(self, text):
        self.set_font(family=self.font, style='B', size=12)
        self.cell(0, 10, text)
        self.ln(14)

    def text_block(self, text, style='', size=11, multi=False):
        self.set_font(family=self.font, style=style, size=size)
        if multi:
            self.multi_cell(0, 6, text)
        else:
            self.cell(0, 10, text)
        self.ln(6)

    def generate_html_table(self, data):
        if len(data) == 0:
            return ''

        columns = list(data[0].keys())
        length = len(columns)

        thead = ''.join(map(
            lambda c: f'<th width="{100 // length}%" align="left">{c}</th>',
            columns,
        ))

        tbody = []
        for item in data:
            row = ''.join(map(lambda x: f'<td>{x}</td>', item.values()))
            tbody.append(f'<tr>{row}</tr>')
        tbody = ''.join(tbody)

        html = '''
        <font face="Arial" size="10">
             <table border="0" width="100%">
                <thead>
                    <tr>{}</tr>
                </thead>
                <tbody>{}</tbody>
            </table>
        </font>
        '''.format(thead, tbody)
        html = html.replace('\n', '')

        return html


class RequestViewSet(viewsets.ModelViewSet):
    serializer_class = RequestSerializer
    pagination_class = StandardResultsSetPagination

    filter_backends = (filters.SearchFilter,)
    search_fields = ('name', 'description', 'user__first_name',
                     'user__last_name',)

    def get_queryset(self,showAll=False):
        libraries_qs = Library.objects.all().only('status', 'sequencing_depth')
        samples_qs = Sample.objects.all().only('status', 'sequencing_depth')
        #   print(libraries_qs.values())


        queryset = Request.objects.select_related('user').prefetch_related(
            Prefetch('libraries', queryset=libraries_qs),
            Prefetch('samples', queryset=samples_qs),
            'files',
        ).order_by('-create_time')

        if not showAll:
            queryset = queryset.filter(sequenced=False)
        if self.request.user.is_staff:
            # Show only those Requests, whose libraries and samples
            # haven't reached status 6 yet
            # TODO: find a way to hide requests
            # queryset = [x for x in queryset if x.statuses.count(6) == 0]
            #queryset = [x for x in queryset if x.statuses.count(5)==0]
            print('staff')
        else:
            queryset = queryset.filter(user=self.request.user)

        #queryset = [x for x in queryset if x.statuses.count(5)==0]



        return queryset

    def list(self, request):
        """ Get the list of requests. """

        showAll = False
        if request.GET.get('showAll') == 'True':
            showAll = True

        queryset = self.filter_queryset(self.get_queryset(showAll))

        try:
            page = self.paginate_queryset(queryset)
        except NotFound:
            page = None

        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)

        return Response(serializer.data)

    def create(self, request):
        """ Create a request. """
        post_data = self._get_post_data(request)
        post_data.update({'user': request.user.pk})
        serializer = self.serializer_class(data=post_data)

        if serializer.is_valid():
            serializer.save()
            return Response({'success': True}, 201)

        else:
            return Response({
                'success': False,
                'message': 'Invalid payload.',
                'errors': serializer.errors,
            }, 400)

    @action(methods=['post'], detail=True)
    def edit(self, request, pk=None):
        """ Update request with a given id. """
        instance = self.get_object()
        post_data = self._get_post_data(request)

        post_data.update({'user': instance.user.pk})

        serializer = self.get_serializer(data=post_data, instance=instance)

        if serializer.is_valid():
            serializer.save()
            return Response({'success': True})

        else:
            return Response({
                'success': False,
                'message': 'Invalid payload.',
                'errors': serializer.errors,
            }, 400)

    @action(methods=['post'],detail=True)
    def mark_as_complete(self,request,pk=None):
        """Mark request as complete, set sequenced to true"""

        instance = Request.objects.filter(pk=pk)



        post_data = self._get_post_data(request)
        override = post_data['override']

        if post_data['override'] == 'False':
            override = False
        else:
            override = True

        def checkifcomplete(element):
            if element==5:
                return True
            else:
                return False


        if override:
            print("Override is true")
            instance.update(sequenced=True)
            return Response({'success':True})

        else:
            print("Override is false")
            #print(instance.statuses)
            #check if all libraries/samples related to this requested have been sequenced
            statuses = [status for x in instance for status  in x.statuses]

            complete = all([checkifcomplete(x) for x in statuses])

            if complete:
                print("all statuses are complete")
                instance.update(sequenced=True)
                return Response({'success':True})
            elif not complete:
                print("there are incomplete statuses")
                return Response({'noncomplete':True})
            else:
                return Response({'error':'error'})


    @action(methods=['post'], detail=True)
    def samples_submitted(self, request, pk=None):
        instance = self.get_object()
        post_data = self._get_post_data(request)
        instance.samples_submitted = post_data['result']
        instance.save(update_fields=['samples_submitted'])
        return Response({'success': True})

    @action(methods=['get'], detail=True)
    def get_records(self, request, pk=None):
        """ Get the list of record's submitted libraries and samples. """
        libraries_qs = Library.objects.all().only(
            'name',
            'barcode',
        )
        samples_qs = Sample.objects.all().only(
            'name',
            'barcode',
            'is_converted',
        )

        instance = Request.objects.filter(pk=pk).prefetch_related(
            Prefetch('libraries', queryset=libraries_qs),
            Prefetch('samples', queryset=samples_qs),
        ).only('libraries', 'samples').first()

        data = [{
            'pk': obj.pk,
            'name': obj.name,
            'barcode': obj.barcode,
            'record_type': obj.__class__.__name__,
            'is_converted': True
            if hasattr(obj, 'is_converted') and obj.is_converted else False,
        } for obj in instance.records]

        data = sorted(data, key=lambda x: x['barcode'][3:])
        return Response(data)

    @action(methods=['get'], detail=True)
    def get_files(self, request, pk=None):
        """ Get the list of attached files for a request with a given id. """
        instance = self.get_object()
        files = instance.files.all().order_by('name')
        serializer = RequestFileSerializer(files, many=True)
        return Response(serializer.data)

    @action(methods=['post'], detail=False,
            authentication_classes=[CsrfExemptSessionAuthentication])
    def upload_files(self, request):
        file_ids = []

        if not any(request.FILES):
            return JsonResponse({
                'success': False,
                'message': 'No files provided.'
            }, status=400)

        for file in request.FILES.getlist('files'):
            f = FileRequest(name=file.name, file=file)
            f.save()
            file_ids.append(f.id)

        return JsonResponse({'success': True, 'fileIds': file_ids})

    @action(methods=['get'], detail=False)
    def get_files_after_upload(self, request):
        file_ids = json.loads(request.query_params.get('file_ids', '[]'))
        error = ''
        data = []

        try:
            files = [f for f in FileRequest.objects.all() if f.id in file_ids]
            data = [
                {
                    'id': file.id,
                    'name': file.name,
                    'size': file.file.size,
                    'path': settings.MEDIA_URL + file.file.name,
                }
                for file in files
            ]

        except Exception as e:
            error = 'Could not get the attached files.'
            logger.exception(e)

        return JsonResponse({
            'success': not error,
            'error': error,
            'data': data,
        })


    @action(methods=['get'],detail=False)
    def download_RELACS_Pellets_Abs_form(self,request):
        print(settings.FILES_PATH)
        file_path = os.path.join(settings.FILES_PATH, 'RELACS_Pellets_Abs_form.xlsx')
        print(file_path)

        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response




    @action(methods=['get'], detail=True)
    def download_deep_sequencing_request(self, request, pk=None):  # pragma: no cover
        """ Generate a deep sequencing request form in PDF. """
        instance = self.get_object()
        user = instance.user
        organization = user.organization.name if user.organization else ''
        cost_unit = instance.cost_unit.name if instance.cost_unit else ''
        objects = list(itertools.chain(
            instance.samples.all(),
            instance.libraries.all(),
        ))
        records = [{
            'name': obj.name,
            'type': obj.__class__.__name__,
            'barcode': obj.barcode,
            'depth': obj.sequencing_depth,
        } for obj in objects]
        records = sorted(records, key=lambda x: x['barcode'][3:])

        pdf = PDF('Deep Sequencing Request')
        pdf.set_draw_color(217, 217, 217)
        pdf.alias_nb_pages()
        pdf.add_page()

        # Deep Sequencing Request info
        pdf.info_row('Request Name', instance.name)
        pdf.info_row('Date', instance.create_time.strftime('%d.%m.%Y'))
        pdf.info_row('User', user.full_name)
        pdf.info_row('Phone', user.phone if user.phone else '')
        pdf.info_row('Email', user.email)
        pdf.info_row('Organization', organization)
        pdf.info_row('Cost Unit', cost_unit)
        pdf.multi_info_row('Description', instance.description)

        y = pdf.get_y()
        pdf.line(pdf.l_margin + 1, y, pdf.fw - pdf.r_margin - 1, y)

        # List of libraries/samples
        heading = 'List of libraries/samples to be submitted for sequencing'
        pdf.set_font('Arial', style='B', size=13)
        pdf.ln(5)
        pdf.cell(0, 10, heading, align='C')
        pdf.ln(10)

        pdf.table_row('#', 'Name', 'Barcode', 'Type',
                      'Sequencing Depth (M)', True)

        for i, record in enumerate(records):
            pdf.table_row(i + 1, record['name'], record['barcode'],
                          record['type'], record['depth'])

        pdf.ln(10)
        y = pdf.get_y()
        pdf.line(pdf.l_margin + 1, y, pdf.fw - pdf.r_margin - 1, y)
        pdf.ln(30)

        # Ensure there is enough space for the signature
        if pdf.get_y() > 265:
            pdf.add_page()
            pdf.ln(20)

        # Signature
        pdf.set_draw_color(0, 0, 0)
        y = pdf.get_y()
        x1_date = pdf.fw / 2
        x2_date = x1_date + 45
        x1_signature = x2_date + 5
        x2_signature = pdf.fw - pdf.r_margin - 1
        pdf.line(x1_date, y, x2_date, y)
        pdf.line(x1_signature, y, x2_signature, y)

        pdf.set_x(x1_date + (x2_date - x1_date) / 2 - 6)
        pdf.cell(12, 10, '(Date)')
        pdf.set_x(x1_signature + 2)
        pdf.cell(0, 10, '(Principal Investigator)')

        pdf = pdf.output(dest='S').encode('latin-1')

        # Generate response
        request_name = normalize(
            'NFKD', instance.name
        ).encode('ASCII', 'ignore').decode('utf-8')
        f_name = request_name + '_Deep_Sequencing_Request.pdf'
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="%s"' % f_name

        return response

    @action(methods=['post'], detail=True,
            authentication_classes=[CsrfExemptSessionAuthentication])
    def upload_deep_sequencing_request(self, request, pk=None):
        """
        Upload a deep sequencing request with the PI's signature and
        change request's libraries' and samples' statuses to 1.
        """
        instance = self.get_object()

        if not any(request.FILES):
            return JsonResponse({
                'success': False,
                'message': 'File is missing.'
            }, status=400)

        instance.deep_seq_request = request.FILES.get('file')
        instance.save()

        file_name = instance.deep_seq_request.name.split('/')[-1]
        file_path = settings.MEDIA_URL + instance.deep_seq_request.name

        instance.libraries.all().update(status=1)
        instance.samples.all().update(status=1)

        return JsonResponse({
             'success': True,
             'name': file_name,
             'path': file_path
        })

    @action(methods=['post'], detail=True, permission_classes=[IsAdminUser])
    def send_email(self, request, pk=None):  # pragma: no cover
        """ Send an email to the user. """
        error = ''

        instance = self.get_object()
        subject = request.data.get('subject', '')
        message = request.data.get('message', '')
        include_failed_records = json.loads(request.POST.get(
            'include_failed_records', 'false'))
        records = []

        # TODO: check if it's possible to send emails at all

        try:
            if subject == '' or message == '':
                raise ValueError('Email subject and/or message is missing.')

            if include_failed_records:
                records = list(instance.libraries.filter(status=-1)) + \
                    list(instance.samples.filter(status=-1))
                records = sorted(records, key=lambda x: x.barcode[3:])

            send_mail(
                subject=subject,
                message='',
                html_message=render_to_string('email.html', {
                    'full_name': instance.user.full_name,
                    'message': message,
                    'records': records,
                }),
                # from_email=settings.SERVER_EMAIL,
                from_email='deepseq@ie-freiburg.mpg.de',
                recipient_list=[instance.user.email],
            )

        except Exception as e:
            error = str(e)
            logger.exception(e)

        return JsonResponse({'success': not error, 'error': error})

    @action(methods=['get'], detail=True)
    def download_complete_report(self, request, pk=None):
        def add_table(document, header, data):
            # Create table
            table = document.add_table(rows=1, cols=len(header))
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(header):
                hdr_cells[i].text = h
            for row in data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)


            # Change font size for all cells
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)

        f_name = 'QC Complete Report.docx'
        response = HttpResponse(content_type='application/vnd.openxmlformats' +
                                '-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{f_name}"'

        instance = self.get_object()
        records = sorted(list(itertools.chain(
            instance.libraries.all(),
            instance.samples.all(),
        )), key=lambda x: x.barcode[3:])

        # Create DOCX document and set default font family
        doc = Document()
        font = doc.styles['Normal'].font
        font.name = 'Arial'

        # Page 1
        doc.add_heading('Complete Report', 0)
        p = doc.add_paragraph('')
        p.add_run('Date, Request ID').bold = True

        doc.add_paragraph('')

        p = doc.add_paragraph('')
        p.add_run('Table of Contents').bold = True
        doc.add_paragraph('Summary')
        doc.add_paragraph('Quality Control of received samples')
        doc.add_paragraph('Library Construction')
        doc.add_paragraph('Cluster Generation and Sequencing')
        doc.add_paragraph('Acknowledgements')
        doc.add_paragraph('Appendix')
        doc.add_page_break()

        # Page 2
        doc.add_heading('General Summary of Workflow', 1)
        doc.add_paragraph()
        doc.add_paragraph('Submitted samples or libraries undergo an ' +
                          'incoming quality control using appropriate ' +
                          'analytical instruments (Fluorometer, Capillary ' +
                          'Electrophoresis, qPCR etc). All samples that ' +
                          'pass international quality standards are ' +
                          'subjected to appropriate library preparation ' +
                          'methods. Qualified libraries are pooled for ' +
                          'multiplex sequencing. An Index Generator ' +
                          'Software assures suitable index design. Pooled ' +
                          'libraries are sequenced to reach desired ' +
                          'depth/coverage using installed sequencing ' +
                          'instruments. Immediately after the sequencing ' +
                          'run bcl to fastq conversion and demultiplexing ' +
                          'is done and the user informed.')
        doc.add_page_break()

        # Page 3
        doc.add_heading('Quality Control of received samples/libraries', 1)
        doc.add_paragraph()
        doc.add_paragraph('All documented measurements were conducted by ' +
                          'the deep sequencing facility, MPI-IE Freiburg. ' +
                          'Raw data and reports of fluorometric ' +
                          'quantification (Qubit) and size distribution ' +
                          'measurements (Fragment Analyzer) can be found as ' +
                          'attachment to each request in Parkour ' +
                          '(parkour.ie-freiburg.mpg.de).')
        header = [
            'Date',
            'ID',
            'Name',
            'L/S',
            'Nuc.Type',
            'ng/µl',
            'bp',
            'Comments',
        ]
        data = []
        for r in records:
            rtype = r.__class__.__name__
            row = [
                r.create_time.strftime('%d.%m.%Y'),
                r.barcode,
                r.name,
                rtype[0],
                r.nucleic_acid_type.name if rtype == 'Sample' else '-',
                r.concentration,
                r.mean_fragment_size if rtype == 'Library' else '-',
                r.comments
            ]
            data.append(row)
        add_table(doc, header, data)
        doc.add_page_break()

        # Page 4
        doc.add_heading('Library Construction', 1)
        doc.add_paragraph()
        doc.add_paragraph('Documentation is only possible if libraries were ' +
                          'constructed in the deep sequencing facility, ' +
                          'MPI-IE Freiburg. Raw data and reports of ' +
                          'fluorometric quantification (Qubit) and size ' +
                          'distribution measurements (Fragment Analyzer) ' +
                          'can be found as attachment to each request in ' +
                          'Parkour (parkour.ie-freiburg.mpg.de). Given ' +
                          'Library Preparation Methods are detailed in the ' +
                          'appendix.')
        lib_prep_objects = LibraryPreparation.objects.filter(
            sample__in=instance.samples.all())
        header = [
            'Date',
            'ID',
            'Name',
            'Protocol',
            'Index I7',
            'Index I5',
            'PCR',
            'ng/µl',
            'bp',
            'nM',
            'Comments',
        ]
        data = []
        for r in lib_prep_objects:
            row = [
                r.create_time.strftime('%d.%m.%Y'),
                r.sample.barcode,
                r.sample.name,
                r.sample.library_protocol.name,
                r.sample.index_i7,
                r.sample.index_i5,
                r.pcr_cycles,
                r.concentration_library,
                r.mean_fragment_size,
                r.nM,
                r.comments
            ]
            data.append(row)
        add_table(doc, header, data)
        doc.add_page_break()

        # Page 5
        doc.add_heading('Cluster Generation and Sequencing', 1)
        doc.add_paragraph()
        header = [
            'Date',
            'ID',
            'Name',
            'Pool ID',
            'Flowcell ID',
            'Sequencer',
            'Depth (M)',
            '% Confident off species reads',
        ]
        data = []
        try:
            flowcell = instance.flowcell.get()
        except Exception:
            flowcell = None
        if flowcell:
            pool_ids = ', '.join(sorted(set(
                flowcell.lanes.values_list('pool__name', flat=True))))
            sequences = flowcell.sequences if flowcell.sequences else []
            conf_reads = {
                s['barcode']: s.get('confident_reads', '')
                for s in sequences
            }
            for r in records:
                row = [
                    flowcell.create_time.strftime('%d.%m.%Y'),
                    r.barcode,
                    r.name,
                    pool_ids,
                    flowcell.flowcell_id,
                    flowcell.sequencer.name,
                    r.sequencing_depth,
                    conf_reads.get(r.barcode, ''),
                ]
                data.append(row)
        add_table(doc, header, data)
        doc.add_page_break()

        # Page 6
        doc.add_heading('Acknowledgements', 1)
        doc.add_paragraph()
        doc.add_paragraph('If data produced in the Deep Sequencing Facility ' +
                          'at MPI-IE, Freiburg is published, include an ' +
                          'acknowledgement in your paper. Also, review if ' +
                          'contributions are substantial and should lead to ' +
                          'an authorship of staff of the facility. ')
        doc.add_paragraph()
        doc.add_paragraph('Additionally, let us know of any publications ' +
                          'involving the facility. Tracking citations and ' +
                          'publications demonstrate the usefulness of the ' +
                          'facility as a research resource which is needed ' +
                          'to obtain further funding.')
        doc.add_paragraph()
        doc.add_paragraph('Example acknowledgement')
        doc.add_paragraph()
        doc.add_paragraph('We thank the Deep Sequencing Facility @ MPI-IE ' +
                          'Freiburg, for performance of quality controls, ' +
                          'library construction and Illumina sequencing.')
        doc.add_page_break()

        # Page 7
        doc.add_heading('Appendix', 1)
        doc.add_paragraph()
        doc.add_paragraph('Detailed list of different library preparation ' +
                          'protocols, sequencing devices and installed ' +
                          'software')

        doc.save(response)
        return response

    def _get_post_data(self, request):
        post_data = {}
        if request.is_ajax():
            post_data = request.data.get('data', {})
            if isinstance(post_data, str):
                post_data = json.loads(post_data)
        else:
            post_data = json.loads(request.data.get('data', '{}'))
        return post_data
